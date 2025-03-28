import re
import json
import logging
import spacy
import PyPDF2
import docx
import os

# Try to load spaCy language model for NER
try:
    nlp = spacy.load("en_core_web_sm")
    HAS_SPACY = True
except:
    HAS_SPACY = False
    logging.warning("spaCy model not available - falling back to basic text processing")

# Try to load ML resume parser
try:
    from utils.ml_parser import MLResumeParser
    ml_parser = MLResumeParser()
    ml_parser.load_skills_classifier()
    HAS_ML_PARSER = True
except:
    HAS_ML_PARSER = False
    logging.warning("ML parser not available - falling back to rule-based parsing")

def extract_resume_data(file_path, file_extension):
    """
    Extract data from resume file
    
    Args:
        file_path (str): Path to the resume file
        file_extension (str): File extension (pdf, docx, json)
    
    Returns:
        dict: Dictionary containing extracted data
    """
    logging.debug(f"Extracting data from {file_path} with extension {file_extension}")
    
    # Extract text based on file type
    if file_extension == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        text = extract_text_from_docx(file_path)
    elif file_extension == 'json':
        return extract_data_from_json(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")
    
    # Extract structured data from text
    return extract_structured_data_from_text(text)

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        if not text.strip():
            logging.warning("PDF text extraction returned empty string, possibly a scanned document")
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {str(e)}")
        raise
    
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract from tables as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {str(e)}")
        raise
    
    return text

def extract_data_from_json(file_path):
    """Extract data from JSON file"""
    try:
        with open(file_path, 'r') as file:
            data = json.load(file)
        
        # Validate required fields
        required_fields = ['skills']
        for field in required_fields:
            if field not in data:
                logging.warning(f"Required field '{field}' missing from JSON resume")
        
        return data
    except Exception as e:
        logging.error(f"Error extracting data from JSON: {str(e)}")
        raise

def extract_structured_data_from_text(text):
    """
    Extract structured data from plain text using NLP
    
    Args:
        text (str): Plain text from resume
    
    Returns:
        dict: Dictionary with structured data
    """
    # Basic structure 
    data = {
        'name': '',
        'email': '',
        'phone': '',
        'skills': [],
        'experience': []
    }
    
    # Use spaCy for NLP processing if available
    if HAS_SPACY:
        doc = nlp(text)
        
        # Extract name
        data['name'] = extract_name(doc, text)
        
        # Extract experience
        data['experience'] = extract_experience(doc, text)
    else:
        # Basic fallback for name extraction
        name_match = re.search(r'^([A-Z][a-z]+ [A-Z][a-z]+)', text.strip(), re.MULTILINE)
        if name_match:
            data['name'] = name_match.group(1)
    
    # Email extraction is consistent
    data['email'] = extract_email(text)
    
    # Phone extraction
    data['phone'] = extract_phone(text)
    
    # Skills extraction
    data['skills'] = extract_skills(doc if HAS_SPACY else None, text)
    
    logging.debug(f"Extracted resume data: {str(data)}")
    return data

def extract_name(doc, text):
    """Extract name from text"""
    # Strategy 1: Try to find the name as the first PERSON entity in the document
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
            
    # Strategy 2: Look for name at the beginning of the resume
    # This assumes the name is often at the top
    lines = text.split('\n')
    for i in range(min(5, len(lines))):
        line = lines[i].strip()
        # Check if line is properly capitalized and doesn't contain too many words
        if line and len(line.split()) <= 4 and all(w[0].isupper() for w in line.split() if w):
            # Avoid mistaking titles or sections for names
            if not any(title in line.lower() for title in ['resume', 'cv', 'curriculum', 'vitae', 'profile']):
                return line
                
    # Strategy 3: Look for pattern of first and last name
    name_pattern = r'([A-Z][a-z]+ [A-Z][a-z]+)'
    name_match = re.search(name_pattern, text[:500])
    if name_match:
        return name_match.group(1)
        
    return ""

def extract_email(text):
    """Extract email from text"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    if email_match:
        return email_match.group(0)
    return ""

def extract_phone(text):
    """Extract phone number from text"""
    # Multiple phone formats
    phone_patterns = [
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # 123-456-7890
        r'\b\(\d{3}\)\s*\d{3}[-.\s]?\d{4}\b',   # (123) 456-7890
        r'\b\d{4}[-.\s]?\d{3}[-.\s]?\d{3}\b',   # International format: 1234 567 890
        r'\b\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b'  # +1-123-456-7890
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            return phone_match.group(0)
    
    return ""

def load_expanded_skills():
    """Load an expanded set of skills"""
    # Common skills in various fields
    expanded_skills = {
        # Programming languages
        "programming": [
            "Python", "Java", "JavaScript", "C++", "C#", "Ruby", "PHP", "Swift", 
            "Kotlin", "Go", "Rust", "TypeScript", "Scala", "R", "MATLAB", "SQL", 
            "Perl", "Bash", "Shell", "Assembly", "Fortran", "COBOL", "HTML", "CSS"
        ],
        # Frameworks and libraries
        "frameworks": [
            "React", "Angular", "Vue.js", "Django", "Flask", "Spring", "Node.js",
            "Express.js", "TensorFlow", "PyTorch", "Keras", "Scikit-learn", "Pandas",
            "NumPy", "jQuery", "Bootstrap", "Laravel", "Ruby on Rails", "ASP.NET",
            "Symfony", ".NET Core", "Xamarin", "Flutter", "SwiftUI"
        ],
        # Databases
        "databases": [
            "MySQL", "PostgreSQL", "SQLite", "MongoDB", "Oracle", "Microsoft SQL Server",
            "Redis", "Cassandra", "Elasticsearch", "DynamoDB", "Firestore", "Neo4j",
            "MariaDB", "CouchDB", "Hadoop", "Hive", "Spark SQL", "IBM Db2"
        ],
        # Tools and platforms
        "tools": [
            "Git", "GitHub", "GitLab", "Docker", "Kubernetes", "AWS", "Azure", "GCP",
            "Jenkins", "Travis CI", "CircleCI", "Ansible", "Terraform", "Puppet", "Chef",
            "Jira", "Confluence", "Slack", "Trello", "Heroku", "Nginx", "Apache",
            "Linux", "Unix", "Windows Server", "macOS", "Eclipse", "Visual Studio",
            "VS Code", "IntelliJ IDEA", "PyCharm", "Sublime Text", "Atom", "Vim"
        ],
        # Data science and analytics
        "data_science": [
            "Machine Learning", "Deep Learning", "AI", "Natural Language Processing",
            "Computer Vision", "Data Mining", "Statistics", "Big Data", "Data Warehousing",
            "ETL", "Business Intelligence", "Tableau", "Power BI", "Looker", "Qlik",
            "SAS", "SPSS", "Data Science", "A/B Testing", "Regression Analysis",
            "Classification", "Clustering", "Decision Trees", "Random Forests", 
            "Neural Networks", "Time Series Analysis", "Predictive Modeling"
        ],
        # Project management and methodologies
        "methodologies": [
            "Agile", "Scrum", "Kanban", "Lean", "Waterfall", "DevOps", "CI/CD",
            "TDD", "BDD", "XP", "ITIL", "Six Sigma", "PMP", "PRINCE2", "SAFe",
            "Product Management", "Requirements Analysis", "UML", "ERD"
        ],
        # Soft skills
        "soft_skills": [
            "Leadership", "Communication", "Problem Solving", "Team Work", "Time Management",
            "Critical Thinking", "Decision Making", "Adaptability", "Creativity", "Collaboration",
            "Presentation", "Negotiation", "Conflict Resolution", "Emotional Intelligence"
        ],
        # Business and management
        "business": [
            "Project Management", "Product Management", "Marketing", "Sales", "CRM",
            "ERP", "Supply Chain", "Logistics", "Finance", "Accounting", "Budgeting",
            "Strategic Planning", "Business Analysis", "Financial Analysis", "Market Research",
            "SWOT Analysis", "Customer Experience", "UX Research", "SEO", "SEM", "Content Marketing"
        ],
        # Design
        "design": [
            "UI/UX", "Graphic Design", "Adobe Photoshop", "Adobe Illustrator", "Adobe XD",
            "Sketch", "Figma", "InDesign", "After Effects", "Premiere Pro", "3D Modeling",
            "Animation", "Wireframing", "Prototyping", "Responsive Design", "Accessibility",
            "Typography", "Color Theory", "Branding", "User Research", "Usability Testing"
        ]
    }
    
    # Flatten the dictionary into a single list of skills
    all_skills = []
    for category, skills in expanded_skills.items():
        all_skills.extend(skills)
    
    return set(all_skills)

def extract_skills(doc, text):
    """Extract skills from text with ML assistance"""
    skills = set()
    
    # Strategy 1: Use ML classifier if available
    if HAS_ML_PARSER:
        # Extract skills section
        skills_section = extract_section(text, 
                                     ["skills", "technical skills", "core competencies", 
                                      "technologies", "qualifications", "expertise"])
        if skills_section:
            # Split by common delimiters and filter out non-skills
            candidates = re.split(r'[,.\n•|\t/&]', skills_section)
            for candidate in candidates:
                skill = candidate.strip()
                if skill and len(skill) > 1:
                    # Use ML model to check if this is a skill
                    if ml_parser.is_skill(skill):
                        skills.add(skill)
    
    # Strategy 2: Use predefined list of common skills
    expanded_skills = load_expanded_skills()
    for skill in expanded_skills:
        # Look for exact matches of multi-word skills
        if len(skill.split()) > 1 and skill.lower() in text.lower():
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text, re.IGNORECASE):
                skills.add(skill)
        # For single-word skills, be more careful to avoid false positives
        elif len(skill.split()) == 1 and len(skill) > 2:
            pattern = r'\b' + re.escape(skill) + r'\b'
            if re.search(pattern, text, re.IGNORECASE):
                skills.add(skill)
    
    # Strategy 3: Use named entities from spaCy if available
    if doc:
        # Find entities that might be skills
        for ent in doc.ents:
            if ent.label_ in ["ORG", "PRODUCT"]:
                # Verify this looks like a skill
                if 2 <= len(ent.text.split()) <= 4 and len(ent.text) > 2:
                    if HAS_ML_PARSER:
                        if ml_parser.is_skill(ent.text):
                            skills.add(ent.text)
                    else:
                        skills.add(ent.text)
    
    # Convert set to list and sort for consistent order
    return sorted(list(skills))

def extract_experience(doc, text):
    """Extract work experience from text with enhanced NLP"""
    experiences = []
    
    # Strategy 1: Extract experience section
    experience_section = extract_section(text, 
                                       ["experience", "work experience", "professional experience", 
                                        "employment history", "work history", "career history"])
    
    if not experience_section:
        return experiences
    
    # Strategy 2: Advanced NLP parsing with spaCy
    if doc:
        # Identify sentences that likely contain job roles or companies
        sentences = [sent.text for sent in doc.sents]
        for sent in sentences:
            # Check if sentence contains indicators of a job description
            if any(word in sent.lower() for word in ["worked", "managed", "developed", "created", 
                                                  "designed", "implemented", "led", "responsible"]):
                if len(sent.split()) > 5:  # Avoid very short fragments
                    experiences.append(sent.strip())
    
    # Strategy 3: Fall back to lines that look like job experience entries
    if not experiences:
        lines = experience_section.split('\n')
        current_exp = ""
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines
            if not line:
                continue
                
            # Check for job title-like patterns or company names
            if (re.match(r'^[A-Z]', line) and 
                any(title in line.lower() for title in ["engineer", "developer", "manager", "director", 
                                                     "analyst", "designer", "consultant", "specialist"])):
                
                # If we were building up previous experience, add it
                if current_exp:
                    experiences.append(current_exp.strip())
                
                # Start new experience entry
                current_exp = line
            else:
                # Continue building current experience
                current_exp += " " + line
        
        # Add the last experience
        if current_exp:
            experiences.append(current_exp.strip())
    
    # If we still don't have experiences, try bullet point extraction
    if not experiences:
        bullets = re.findall(r'•(.*?)(?=•|\n\n|$)', experience_section, re.DOTALL)
        experiences = [b.strip() for b in bullets if b.strip()]
    
    return experiences

def extract_section(text, section_names):
    """
    Extract text from a specific section of the resume
    
    Args:
        text (str): Full resume text
        section_names (list): Possible names for the section
    
    Returns:
        str: Text from the section or empty string if not found
    """
    text = text.lower()
    
    section_patterns = []
    for name in section_names:
        # Create patterns that match section headers
        # Examples: "SKILLS", "Skills:", "TECHNICAL SKILLS", etc.
        patterns = [
            rf'(?i)\b{re.escape(name)}\s*:',  # "Skills:"
            rf'(?i)\b{re.escape(name)}\b',     # "Skills" as whole word
            rf'(?i)\b{re.escape(name.upper())}\b',  # "SKILLS"
            rf'(?i)\b{re.escape(name.title())}\b'   # "Skills"
        ]
        section_patterns.extend(patterns)
    
    # Find the start of the specified section
    section_start = -1
    section_name_found = ""
    
    for pattern in section_patterns:
        match = re.search(pattern, text)
        if match and (section_start == -1 or match.start() < section_start):
            section_start = match.start()
            section_name_found = match.group(0)
    
    if section_start == -1:
        return ""
    
    # Find the start of the next section
    next_section_start = len(text)
    common_sections = [
        "education", "experience", "work experience", "skills", "projects", 
        "certifications", "awards", "publications", "languages", "interests",
        "summary", "objective", "profile", "contact", "references", "volunteer",
        "extracurricular", "technical skills", "professional experience"
    ]
    
    # Exclude the section we're currently extracting
    for name in section_names:
        if name in common_sections:
            common_sections.remove(name)
    
    # Create patterns for the next section
    next_section_patterns = []
    for name in common_sections:
        patterns = [
            rf'(?i)\b{re.escape(name)}\s*:',  # "Education:"
            rf'(?i)\b{re.escape(name.upper())}\b',  # "EDUCATION"
            rf'(?i)\b{re.escape(name.title())}\b'   # "Education"
        ]
        next_section_patterns.extend(patterns)
    
    # Find the closest next section
    for pattern in next_section_patterns:
        match = re.search(pattern, text[section_start + len(section_name_found):])
        if match:
            candidate_start = section_start + len(section_name_found) + match.start()
            if candidate_start < next_section_start:
                next_section_start = candidate_start
    
    # Extract the section text
    section_text = text[section_start + len(section_name_found):next_section_start]
    
    # Clean up the text
    section_text = section_text.strip()
    
    return section_text